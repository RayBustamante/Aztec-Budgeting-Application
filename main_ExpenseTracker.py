import tkinter as tk
from tkinter import messagebox, scrolledtext, filedialog
import subprocess
import sys
import csv
import os
from  datetime import datetime
import pandas as pd
import time
import speech_recognition as sr
import numpy as np
from openai import OpenAI


from docx import Document

# module for detecting numbers
import re

# Machine Learning imports
import pandas as pd
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.model_selection import train_test_split
from sklearn.svm import SVC
from sklearn.metrics import classification_report

import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk


#-----------------------------------------------------------------
# by Ramon C Bustamante

# Using the inputs provided, this program will try to keep track of all user expenses.
# Which will be downloaded into several csv files containing purchases.
# Additionally, a log file containing the chatbox will be saved to a local txt file that will display
# whenver new purchases are input.
# Lastly, it will analyize the data provided with those of your peers and based on the comparison,
# you may be able to save a document file containing all of your information along with A.I generative
# financial advice tailored to your situation using openAI API's.
#-----------------------------------------------------------------

wantToStartOver = False
expenseLimit = ""
currentSpent = None
aboveAverage = None
monthlyIncome = 0


# Create Personal Information Window
#-----------------------------------------------------------------
root = tk.Tk()
root.title("Account Form")
root.geometry("1274x714")

# Create Background Image
#-----------------------------------------------------------------
background_image2 = tk.PhotoImage(file="Aztec_Menu_Background_Empty.png")
background_label = tk.Label(root, image=background_image2)
background_label.place(relwidth=1, relheight=1)

background_image_button = tk.PhotoImage(file="button_background_red.png")

# Functions
#-----------------------------------------------------------------

def machineLearning_Categorization(input_text):
    
    # A.I Machine Learning model for differenciating text input based on a csv file with "Descrption" & "Category"
    # May need additional datasets for more accurate prediction...
    # At mininum, ensure 300 dataset rows containining descriptions and category
    
    # Support Vector Machine Model 
    data = pd.read_csv("machineLearning_Dataset.csv")

    X = data['description']
    y = data['category']

    print("Machine Learning Model Running...")
    # Split data into train and test sets
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42)

    # Vectorize text data using TF-IDF
    vectorizer = TfidfVectorizer()
    X_train_vectorized = vectorizer.fit_transform(X_train)
    X_test_vectorized = vectorizer.transform(X_test)

    # Train Support Vector Machine (SVM) classifier
    classifier = SVC(kernel='linear')
    classifier.fit(X_train_vectorized, y_train)

    # Evaluate classifier
    y_pred = classifier.predict(X_test_vectorized)
    print(classification_report(y_test, y_pred))

    # Runs the model
    new_expense = [str(input_text)]
    new_expense_vectorized = vectorizer.transform(new_expense)
    predicted_category = str(classifier.predict(new_expense_vectorized).item())
    
    print("Machine Learning Model Finished...")
    
    return predicted_category
    

def type_output(input_text2, chatboxExpense, index=0):
    # Outputs to chatboxx and make it type slowly like in chatGPT
    if index < len(input_text2):
        chatboxExpense.insert(tk.END, input_text2[index])
        index += 1
        chatboxExpense.after(20, type_output, input_text2, chatboxExpense, index)
        chatboxExpense.yview_moveto(1.0)
        saveChatBoxToFile()
        
def type_outputNoSave(input_text2, chatboxExpense, index=0):
    # Outputs to chatboxx and make it type slowly like in chatGPT
    if index < len(input_text2):
        chatboxExpense.insert(tk.END, input_text2[index])
        index += 1
        chatboxExpense.after(20, type_output, input_text2, chatboxExpense, index)
        chatboxExpense.yview_moveto(1.0)
        
def type_outputFaster(input_text2, chatboxExpense, index=0):
    # Outputs to chatboxx and make it type slowly like in chatGPT
    if index < len(input_text2):
        chatboxExpense.insert(tk.END, input_text2[index])
        index += 1
        chatboxExpense.after(100, type_output, input_text2, chatboxExpense, index)
        chatboxExpense.yview_moveto(1.0)

def start_OverQuestion():
    
    #Ask if you really want to start over if a file already exist. Doing so will clear all local files
    # to ensure clean and complete slate for new inputs and reduce the chance of corrupted files.
    file_path = "chatbox_content.txt"
    if wantToStartOver == False and os.path.exists(file_path):
        #check
        result = messagebox.askyesno("Load Files Found!", "You currently have a pre-existing log file. Would you like to load file?")
        if result:
            #If yes load file
            loadExpenseLogFiles()           
        else:
            #If no then return but also clear out data from all local expense file.
            result = messagebox.askyesno("Are you sure?", "Are you sure you want to start over? This will delete your current log files.")
            if result:
                #If yes load file
                deleteExpenseDataFiles()
                return
                loadExpenseLogFiles()           
            else:
                #If no then return but also clear out data from all local expense file.
                loadExpenseLogFiles() 
def deleteExpenseDataFiles():
    
    #Delete all local files for each category
    
    if os.path.exists("housingExpense.csv"):
        os.remove("housingExpense.csv")
        print("Deleting Local Files")
        return
    elif os.path.exists("foodExpense.csv"):
        os.remove("foodExpense.csv")
        print("Deleting Local Files")
        return
    elif os.path.exists("transportationExpense.csv"):
        os.remove("transportationExpense.csv")
        print("Deleting Local Files")
        return
    elif os.path.exists("schoolsuppliesExpense.csv"):
        os.remove("schoolsuppliesExpense.csv")
        print("Deleting Local Files")
        return
    elif os.path.exists("entertainmentExpense.csv"):
        os.remove("entertainmentExpense.csv")
        print("Deleting Local Files")
        return
    elif os.path.exists("personalcareExpense.csv"):
        os.remove("personalcareExpense.csv")
        print("Deleting Local Files")
        return
    elif os.path.exists("technologyExpense.csv"):
        os.remove("technologyExpense.csv")
        print("Deleting Local Files")
        return
    elif os.path.exists("miscellaneousExpense.csv"):
        os.remove("miscellaneousExpense.csv")
        print("Deleting Local Files")
        return



def containsLetterAndNumber(input_string):
    
    # function to ensure each input has amount and description of purchase
    has_letter = False
    has_number = False
    
    # Split the input string by space
    parts = input_string.split()
    
    # Check each part of the string
    for part in parts:
        # Check if the part contains letters and numbers
        for char in part:
            if char.isalpha():
                has_letter = True
            elif char.isdigit():
                has_number = True
    
    # Return True if both a letter and a number are found
    return has_letter and has_number

def grabFirstNumericValue(input_text):
    # function that will grabs the first numeric value of an input for $ purchase.
    match = re.search(r'\d+(\.\d+)?', input_text)
    
    # If a match is found, return the numeric value
    if match:
        return float(match.group())
    else:
        return None  # Return None if no numeric value is found
    

def handle_enter(event):
    # function to catch enter key pressed for inputs.
    
    input_text = entry_Expense.get()
    write_Chatbox(input_text)
    
def write_Chatbox(input_text):
    
    # writes to chatbox the data and content of the input_text.
    global wantToStartOver
    
    # makes a check to make sure if user realy really wants to start over.
    if wantToStartOver == False:
        start_OverQuestion()
        wantToStartOver = True
        #cancels the input until this is over with
        return
    
    if '#' in input_text:
            clean_string = input_text.replace("#", "")
            input_text = ("     Notes: " + str(clean_string) )
            saveChatBoxToFile()
            if input_text:
                chatboxExpense.insert(tk.END, "\n")
                entry_Expense.delete(0, tk.END)  # Clear the entry box after sending the input
                type_output(input_text, chatboxExpense)
        
    if containsLetterAndNumber(input_text):     
        #IF not then do this...
            
        # grabs only the numeric values
        
        dollarAmount = float(grabFirstNumericValue(input_text))
        formatted_dollarAmount = format_Dollar(dollarAmount)
        # Run the machine learning model
        expenseCategory = str(machineLearning_Categorization(input_text))
        #checking...
        #print(expenseCategory)
            
        #get timestamp
        current_datetime = datetime.now()
            
        current_day = current_datetime.day
        month_name = current_datetime.strftime("%B")
        current_time = current_datetime.strftime("%I:%M %p")
            
        dateString = (str(month_name) + " " + str(current_day) + ", " + str(current_time) )
            
        input_text2 = ("You have spent " + str(formatted_dollarAmount) + " on your ['" + str(expenseCategory) +
                            "'] expenses. \n     User Input: " + str(input_text) + "\n     Date: " + str(dateString))
            
            #save chatbox
            
        dateMonthDayYear = current_datetime.strftime("%m/%d/%Y %I:%M:%S %p")
        loadSaveExpensesDate(dollarAmount, expenseCategory,dateMonthDayYear)
            
            
        #print(numbers)
        if input_text:
            chatboxExpense.insert(tk.END, "\n * ")
            entry_Expense.delete(0, tk.END)  # Clear the entry box after sending the input
            type_output(input_text2, chatboxExpense)
            
            
                
    else:
        messagebox.showerror("Error", "Invalid expense input. Must contain a numerical amount and a description. \n\nEntered Input: " + input_text)

def loadExpenseLogFiles():

    #function for loading log files and checks if one exist
    global wantToStartOver
    wantToStartOver = True
    file_path = "chatbox_content.txt"

    # Check if the file exists
    if os.path.exists(file_path):
        with open(file_path, "r") as file:
            content = file.read()
            chatboxExpense.delete(1.0, tk.END)
            type_outputFaster(content, chatboxExpense)
    else:
        messagebox.showerror("File Not Found", "There are currently no Expense Log Files. Please start typing to create a new one")


def saveChatBoxToFile():
    # saves the chatbox content for future use in a local txt file.
    chatbox_content = chatboxExpense.get("1.0", tk.END)
    
    file_path = "chatbox_content.txt"

    # Check if the file already exists
    if not os.path.exists(file_path):
        
        with open(file_path, "w") as file:
            
            file.write(chatbox_content)
            
        #messagebox.showinfo("No Log Found", "Creating a new Log File")

    else:
        with open(file_path, "w") as file:
            file.write(chatbox_content)
        #messagebox.showinfo("File Saved", "Your content has been saved to Log File")             

        
def getExpenseLimit(expenseCategory):
    
    #retrieve expense limit set out from Expense page. 
    
    global monthlyIncome
    filename = "studentUser_newFile.csv"
    userData = []
    expenseLimit = 0
    
    with open(filename, 'r') as file:
        reader = csv.reader(file)
        for row in reader:
            userData.append(row)   
            
    if expenseCategory == "Housing":
        expenseLimit = userData[1][8]
    elif expenseCategory == "Food":
        expenseLimit = userData[1][9]
    elif expenseCategory == "Transportation":
        expenseLimit = userData[1][10]
    elif expenseCategory == "School Supplies":
        expenseLimit = userData[1][11]
    elif expenseCategory == "Entertainment":
        expenseLimit = userData[1][12]
    elif expenseCategory == "Personal Care":
        expenseLimit = userData[1][13]
    elif expenseCategory == "Technology":
        expenseLimit = userData[1][14]
    else:
        expenseLimit = userData[1][16]
    
    monthlyIncome = userData[1][5]
    
    return expenseLimit
               
def loadSaveExpensesDate(dollarAmount, expenseCategory,dateMonthDayYear):
    #load save expense data from local csv file.
    
    global label_Graph_ExpenseTitle
    global expenseLimit
    global currentCategory
    global currentSpent
    
    fileName = ""
    
    if expenseCategory == "Housing":
        filename = "housingExpense.csv"
    elif expenseCategory == "Food":
        filename = "foodExpense.csv"
    elif expenseCategory == "Transportation":
        filename = "transportationExpense.csv"
    elif expenseCategory == "School Supplies":
        filename = "schoolsuppliesExpense.csv"
    elif expenseCategory == "Entertainment":
        filename = "entertainmentExpense.csv"
    elif expenseCategory == "Personal Care":
        filename = "personalcareExpense.csv"
    elif expenseCategory == "Technology":
        filename = "technologyExpense.csv"
    else:
        filename = "miscellaneousExpense.csv"
       
    currentCategory = expenseCategory
    
    #list needed for graph
    chartXAxisDate = []
    chartYAxisExpense = []
    
    if not os.path.exists(filename):
        # Create the CSV file with headers
        with open(filename, "w", newline='') as file:
            writer = csv.writer(file)
            writer.writerow(["Date", "Expense"])
    
    # Read existing data from the CSV file into a list
    expenseDataFile = []

    with open(filename, "r", newline='') as file:
        reader = csv.reader(file)
        next(reader)  # Skip the header row
        expenseDataFile.extend(row for row in reader)
        
    expenseDataFile.append([dateMonthDayYear, dollarAmount])

    # Append and then create graphs
    chartXAxisDate = [str(row[0]) for row in expenseDataFile]
    chartYAxisExpense = [float(row[1]) for row in expenseDataFile]

    # Write the updated data back to the CSV file
    with open(filename, "w", newline='') as file:
        writer = csv.writer(file)
        writer.writerow(["Date", "Expense"])
        writer.writerows(expenseDataFile)

    create_area_chart(chartXAxisDate, chartYAxisExpense, expenseCategory)
    
    if label_Graph_ExpenseTitle.winfo_exists():
        label_Graph_ExpenseTitle.destroy()
        
    # Create the Labels for Calculation
    label_Graph_ExpenseTitle = tk.Label(root, text=" ", font=("Times New Roman", 11), bg='#9A162B')
    label_Graph_ExpenseTitle.grid(row=1, column=1, padx=(150, 5), pady=(5, 5), sticky="w")
    
    # Iterate through YAxis which contained $ expenses
    total = 0
    for num in chartYAxisExpense:
        total = total + float(num)
    
    currentSpent = total
       
    # so I add them together and then create a string value that I will send through start_typing() function.   
    avaliableBalance = float(getExpenseLimit(expenseCategory)) - float(total)

    formatted_Limit = format_Dollar(int(getExpenseLimit(expenseCategory)))
    formatted_Total = "${:.2f}".format(total)
    formatted_Balance = "${:.2f}".format(avaliableBalance)
    
    expenseLimit = ("Monthly " + str(expenseCategory) + " Expense: " + str(formatted_Limit) +
                                                   "\nSpent: " + str(formatted_Total) +
                                                   "\nAvaliable Balance: " + str(formatted_Balance))
    
    start_typing(expenseLimit, label_Graph_ExpenseTitle)

def format_Dollar(value):
    # Format the float value as a dollar amount
    dollar_amount = "${:.2f}".format(value)
    
    # Add a leading zero if the value is less than 1
    if value < 1:
        dollar_amount = "0" + dollar_amount
    
    return dollar_amount
  
def start_typing(expenseLimit, label_Graph_ExpenseTitle):
    
    label_Graph_ExpenseTitle.config(text="")
    type_to_label(expenseLimit, label_Graph_ExpenseTitle)

def type_to_label(input_text, label, index=0):
    
    # Outputs to label and makes it type slowly
    if index < len(input_text):
        label.config(text=label.cget("text") + input_text[index], fg="#D0C395") 
        index += 1
        label.after(100, type_to_label, input_text, label, index)  


def create_area_chart(chartXAxisDate, chartYAxisExpense, expenseCategory):
    #create a plot chart to visually display current spending habits.
    
    global table_frame
    
    for widget in table_frame.winfo_children():
        widget.destroy()

    table_frame.config(bg="white")
    fig = plt.Figure(figsize=(6.7, 4), dpi=100)
    ax = fig.add_subplot(111)
    
    # MAKE SURE THE Y AXIS IS A NUMERIC VALUE!!!!! Or else it won't sort!!
    x = chartXAxisDate
    y = chartYAxisExpense

    # Plot the area chart
    ax.fill_between(x, y, color="skyblue", alpha=0.4)

    # Set titles for x and y axis
    ax.set_xlabel('Dates')
    ax.set_ylabel('Expediture')

    # Rotate x-axis labels
    ax.set_xticklabels(x, rotation=90)

    # Create a canvas for the Figure
    canvas = FigureCanvasTkAgg(fig, master=table_frame)
    canvas.draw()
    canvas.get_tk_widget().pack()

    # Add navigation toolbar for zooming
    toolbar = NavigationToolbar2Tk(canvas, table_frame)
    toolbar.update()
    canvas.get_tk_widget().pack()
    
    button_SaveDoc = tk.Button(root, text="Analyze", font=("Times New Roman", 15),fg="#D0C395", borderwidth=0, width=130, height=35, compound="center", image=background_image_button, relief="flat", command=analysize_Data)
    button_SaveDoc.place(x=730, y=650)
    
def button_Back():
    # return to menu page.
    
    try:
      subprocess.Popen([sys.executable, "main_newFile_Menu.py"])
      sys.exit()
    except Exception as e:
      print("An error occurred:", e)
      
def record_Speech():
    # Function to record speech
    def start_Recording():
        nonlocal text_labelCountdown
        global text_labelCountdown2
        
        # Initialize the recognizer
        recognizer = sr.Recognizer()

        # Capture audio from the microphone for 3 seconds
        with sr.Microphone() as source:
            text_labelCountdown.config(text="Listening for 5 seconds...")
            audio = recognizer.listen(source, timeout=5, phrase_time_limit=3)

        # Recognize speech using Google Speech Recognition
        try:
            text_labelCountdown.config(text="Recognizing...")
            text = recognizer.recognize_google(audio)
            text_labelCountdown.config(text="Voice Capture: \n\n '" + text + "' " + " \n\nSending voice capture to program")
            root.after(3000, write_Chatbox, text)
            
            
            
        except sr.UnknownValueError:
            text_labelCountdown.config(text="Sorry, could not understand audio.")
        except sr.RequestError as e:
            text_labelCountdown.config(text="Error occurred: {0}".format(e))
        
        record_window.after(3000, record_window.destroy)

    # Function to start recording after a delay
    def delayed_StartRecording():
        # Wait for 3 seconds before starting recording
        root.after(100, start_Recording)

    def type_CountdownLabel(timer, label):
        # Display the timer value
        label.config(text="Recording in " + str(timer) + "...")
        if timer > 0:
            # Schedule the next countdown after a brief delay
            label.after(1000, type_CountdownLabel, timer - 1, label)
        else:
            # Start the speech recording when the countdown reaches 0
            label.config(text="Recording now...")
            delayed_StartRecording()

    def start_Countdown(label):
        # Start the countdown from 3
        type_CountdownLabel(3, label)
    
    # Create a new window for speech recording
    record_window = tk.Toplevel(root)
    record_window.geometry("335x264")
    record_window.title("Speech Recording")
    
    # Load the background image
    background_image2_voice = tk.PhotoImage(file="voiceActivation_Background.png")
    background_label_voice = tk.Label(record_window, image=background_image2)
    background_label_voice.place(relwidth=1, relheight=1)
    
    text_labelCountdown = tk.Label(record_window, text="", font=("Arial", 12), bg='#9A162B', fg='white')
    text_labelCountdown.pack(padx=10, pady=(50, 5))
    text_labelCountdown2 = tk.Label(record_window, text="", font=("Arial", 12), bg='#9A162B', fg='white')
    text_labelCountdown2.pack(padx=10, pady=50)

    start_Countdown(text_labelCountdown)
 
def chat_GPT(prompt):
    
    # Create a prompt to send current expense expense data to chatGPT and retrieve viable financial advice
    # tailored to the user's current expense rate.
    
    #Due to security reasons, I removed my API key and now it requires the user entry instead of a predefined key.
    
    apiInput = str(input("Enter API key"))
    
    try:
        OPENAI_API_KEY = apiInput
        client = OpenAI(api_key=OPENAI_API_KEY)
        
        completion = client.chat.completions.create(
          model="gpt-3.5-turbo",
          messages=[
            {"role": "system", "content": "You are an financial advisor that will give a student some advice on his or her's expenses and spending habits"},
            {"role": "user", "content": prompt}
          ]
        )

        print(completion.choices[0].message.content)
        
        advice = (completion.choices[0].message.content)
        
        return advice
    
    except:
        messagebox.showerror("Error API request", "If you recieved this error, the API key provided is invalid.")
        return False

        
    
def analysize_Data():
    #compare the data the user inputs with data from other students.
       
    global currentCategory
    global table_frame
    global currentSpent
    global chatboxExpense
    global aboveAverage
    global monthlyIncome
    
    print(currentCategory)
    expenseCategory = currentCategory
    
    studentData = []
    index = 0
     
    if expenseCategory == "Housing":
        index = 7
    elif expenseCategory == "Food":
        index = 8
    elif expenseCategory == "Transportation":
        index = 9
    elif expenseCategory == "School Supplies":
        index = 10
    elif expenseCategory == "Entertainment":
        index = 11
    elif expenseCategory == "Personal Care":
        index = 12
    elif expenseCategory == "Technology":
        index = 13
    else:
        index = 15
        
    #start pandas data analysis
    expenses_df = pd.read_csv('student_spending.csv')

    # find specific expense based on current calculation
    column_stats = expenses_df.iloc[:, index].agg(['max', 'min', 'mean'])

    #initialize values
    highest_value = 0
    lowest_value = 0
    average_value = 0


    highest_value = int(column_stats['max'])
    lowest_value = int(column_stats['min'])
    average_value = int(column_stats['mean'])
    
    currentSpent = int(currentSpent)
    
    #grab expense category value 
    dataValue = expenses_df.iloc[:300, index]
    yAxisValue = dataValue.tolist()
    
    #grab monthly income
    dataValue = expenses_df.iloc[:300, 5]
    xAxisValue = dataValue.tolist()
    
    
    for widget in table_frame.winfo_children():
        widget.destroy()
        
    # linear regression
    slope, intercept = np.polyfit(xAxisValue, yAxisValue, 1)
    line_of_best_fit = slope * np.array(xAxisValue) + intercept

    # Create a Figure and a set of subplots
    fig, ax = plt.subplots(figsize=(6, 4))

    # Plot the scatterplot
    ax.scatter(xAxisValue, yAxisValue, color='blue', alpha=0.7, label='Data')
    ax.plot(xAxisValue, line_of_best_fit, color='red', label='Line of Best Fit (linear regression)')

    # Set labels and title
    ax.set_xlabel('Monthly Income')
    ax.set_ylabel('Expense Category')
    ax.set_title('Scatterplot of ' + str(currentCategory) + " expense")
    ax.legend()

    # Create a canvas for the Figure
    canvas = FigureCanvasTkAgg(fig, master=table_frame)
    canvas.draw()
    canvas.get_tk_widget().pack()

    # Add navigation toolbar for zooming
    toolbar = NavigationToolbar2Tk(canvas, table_frame)
    toolbar.update()
    toolbar.pack()    
    
    
    percentage_difference = ((currentSpent - average_value) / average_value) * 100
    
    formatted_percentage = "{:.2f}".format(abs(percentage_difference))
    
    date_formatted = current_datetime.strftime("%m/%d/%Y")    
    
    string = ""
    if average_value < currentSpent: #if the user expense is greater than the average
        aboveAverage = True
        string = ("\n\nAnalysis Date: " + str(date_formatted) + "\n-----------------\nTotal Entries of Students: 1000 entries \n\nHighest Expenses: " + str(highest_value)
                  + "\nLowest Expenses: " + str(lowest_value) + "\nAverage Expense: " + str(average_value)
                  + "\nExpense Category: " + str(currentCategory)
                  + "\n\nUser Monthly Income: " + str(monthlyIncome)
                  + "\nCurrent Expense : " + str(currentSpent)
                  + "\nStatus: Exceeds average by a difference of " + str(formatted_percentage)
                  + "%\n\nData analysis complete...")
    elif average_value > currentSpent: #if the user expense is lower than the average
        aboveAverage = False
        string = ("\n\nAnalysis Date: " + str(date_formatted) + "\n-----------------\nTotal Entries of Students: 1000 entries \n\nHighest Expenses: " + str(highest_value)
                  + "\nLowest Expenses: " + str(lowest_value) + "\nAverage Expense: " + str(average_value)
                  + "\nExpense Category: " + str(currentCategory)
                  + "\n\nUser Monthly Income: " + str(monthlyIncome)
                  + "\nCurrent Expense : " + str(currentSpent)
                  + "\nStatus: Below average by a difference of " + str(formatted_percentage)
                  + "%\n\nData analysis complete...")
        
    type_outputNoSave(string, chatboxExpense)
    
    save_File()
 
def save_File():
    global aboveAverage
    
    #save files as word doc
    global expenseLimit
    
    
    # Give user option to save as Word doc
    result = messagebox.askyesno("Save as Word Document?", "Would you like to save the contents of your log files into a word document?")
    # Include an option to add financial advice at the end of the doc. I made it an option cause
    # they are charging me for the service that I paid for (Ramon B.)
    if result:
        
        result = messagebox.askyesno("Would you like A.I Advice?", "Would you like to include financial advices tailored to your current situation A.I at the end of your document? \nThis will require an user input of a valid API key in the terminal entry")
        if result:
            #retrieve the text from chatbox
            chat_text = chatboxExpense.get("1.0", tk.END)

            date_formatted = current_datetime.strftime("%m-%d-%Y")
            
            #initial file names
            initialFileName = "Aztec Budget Logfile " + str(date_formatted)
            # Ask the user to choose a file location for saving the Word document
            filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=initialFileName)
            
            #If there is a file path...
            if filepath:
                doc = Document()
                
                # Add the text from the chatbox to the Word document
                doc.add_paragraph(chat_text)
                savings = 0
                
                
                # check to see if file exist
                if os.path.isfile("studentUser_newFile.csv"):
                    
                    #read file and import savings value                    
                    data = []
                    with open("studentUser_newFile.csv", mode='r', newline='') as file:
                        reader = csv.reader(file)
                        for row in reader:
                            data.append(row)
                    savings = data[1][18]
                else:
                    messagebox.showerror("Error", "No account has been created yet. Please create an account using the New File button.")
                
                prompt = ""
                
                if aboveAverage == True:
                    prompt = ("My expenses are above the average of other students. Based on my expenses. What are some financial advice you can give me? Never mention my name" + str(expenseLimit) )
                elif aboveAverage == False:
                    prompt = ("My expenses are below the average of other students. Based on my expenses. What are some financial advice you can give me? Never mention my name" + str(expenseLimit) )
                    
                prompt = ("Based on my expenses. What are some financial advice you can give me?" + str(expenseLimit) )                
                # call chatGPT and get some financial advice
                financialAdvice = chat_GPT(prompt)
                
                if financialAdvice == False:
                    messagebox.showerror("Error", "Saving file without Financial advice.")
                    #Save as plain if failed authentication
                    chat_text = chatboxExpense.get("1.0", tk.END)

                    date_formatted = current_datetime.strftime("%m-%d-%Y")
                    
                    initialFileName = "Aztec Budget Logfile " + str(date_formatted)
                    # Ask the user to choose a file location for saving the Word document
                    filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=initialFileName)

                    if filepath:
                        doc = Document()
                        
                        # Add the text from the chatbox to the Word document
                        doc.add_paragraph(chat_text)
                        
                        #add into the document
                        expenseLimit = (str(expenseLimit) + "\nDate: " + str(date_formatted))
                        doc.add_paragraph(expenseLimit)
                        
                        # Save the Word document to the chosen file location
                        doc.save(filepath)
                    
                    
                else:
                    #add into the document
                    expenseLimit = (str(expenseLimit) + "\nDate of Submission: " + str(date_formatted) + "\n\n---Expense Advice---")
                    doc.add_paragraph(expenseLimit)
                    doc.add_paragraph(financialAdvice)
                    
                    prompt = ("Also, with this amount of savings, should I consider investing my money or saving it? Don't mention my name when you write. Savings: " + str(savings) )                
                    # call chatGPT and get some financial advice
                    financialAdvice = chat_GPT(prompt)
                    
                    #add into the document
                    expenseLimit = ""
                    expenseLimit = (str(expenseLimit) + "\nDate of Submission: " + str(date_formatted) + "\n\n---Savings Advice---")
                    doc.add_paragraph(expenseLimit)
                    doc.add_paragraph(financialAdvice)                
                    
                    # Save the Word document to the chosen file location
                    doc.save(filepath)  
        else:
            
            # Plain word document without the A.I Advice.
            chat_text = chatboxExpense.get("1.0", tk.END)

            date_formatted = current_datetime.strftime("%m-%d-%Y")
            
            initialFileName = "Aztec Budget Logfile " + str(date_formatted)
            # Ask the user to choose a file location for saving the Word document
            filepath = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")], initialfile=initialFileName)

            if filepath:
                doc = Document()
                
                # Add the text from the chatbox to the Word document
                doc.add_paragraph(chat_text)
                
                #add into the document
                expenseLimit = (str(expenseLimit) + "\nDate: " + str(date_formatted))
                doc.add_paragraph(expenseLimit)
                
                # Save the Word document to the chosen file location
                doc.save(filepath)  
    else:
        return

# Entries and Labels
#-----------------------------------------------------------------
expensesTitle_label = tk.Label(root, text="Expense Tracker Logs", font=("Times New Roman", 28),fg="#D0C395", bg='#9A162B')
expensesTitle_label.grid(row=0, column=0, padx=(150, 5), pady=(50, 10), sticky="w")

expensesTitle_label2 = tk.Label(root, text="Type and enter all the daily expenses with their amount and source. \nThis tool will track of all your expenses using seamless Machine Learning modules \n to categorize and catalogie them to a local file. It also allows for note taking", font=("Times New Roman", 10),fg="#D0C395", bg='#9A162B')
expensesTitle_label2.grid(row=1, column=0, padx=(110, 5), pady=(10, 10))

# Chatbox
#-----------------------------------------------------------------
current_datetime = datetime.now()
date_formatted = current_datetime.strftime("%m/%d/%Y")

current_date = "Log Date: " + str(date_formatted)

# this simply stops the chatbox from being editable
def ignore_click(event):
  return "break"

chatboxExpense = scrolledtext.ScrolledText(root, width=50, height=20)
chatboxExpense.grid(row=2, column=0, padx=(120,15), pady=10, sticky="w")

chatboxExpense.configure(bg='white', insertbackground='white')
chatboxExpense.configure(font=("Times New Roman", 12))

chatboxExpense.bind("<Button-1>", ignore_click)

initial_text = "Welcome to Expense Tracker!\nStart entering your expenses below. \n\n" + str(current_date)


chatboxExpense.insert(tk.END, initial_text)

# Expense Input
#-----------------------------------------------------------------
label_Expense = tk.Label(root, text="Enter your expenses (i.e. $50 Starbucks.) \n If you want to add notes (i.e.# Here is a Note)",fg="#D0C395", font=("Times New Roman", 10), bg='#9A162B', anchor="w")
label_Expense.place(x=180, y=600)

entry_Expense = tk.Entry(root, width=50)
entry_Expense.place(x=180, y=650)

entry_Expense.bind("<Return>", handle_enter)

# Graph and Labels
#-----------------------------------------------------------------

# Create frame for the table
table_frame = tk.Frame(root, bg="#9A162B")
table_frame.grid(row=2, column=1, padx=20, pady=2, sticky="nsew")


# Load Log Files
#-----------------------------------------------------------------
label_Graph_ExpenseTitle = tk.Label(root, text="", font=("Times New Roman", 11),fg="#D0C395", bg='#9A162B')
label_Graph_ExpenseTitle.place(x=820, y=60)

# Extra manual load button.
button_Back = tk.Button(root, text="Back", font=("Times New Roman", 15),fg="#D0C395", borderwidth=0, width=100, height=25, compound="center", image=background_image_button, relief="flat", command=button_Back)
button_Back.place(x=25, y=15)

# Voice Activation Input

button_Voice = tk.Button(root, text="Voice Input", font=("Times New Roman", 15),fg="#D0C395", borderwidth=0, width=130, height=35, compound="center", image=background_image_button, relief="flat", command=record_Speech)
button_Voice.place(x=530, y=650)

# Start the main event loop
root.mainloop()